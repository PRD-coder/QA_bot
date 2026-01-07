import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document
import requests
from bs4 import BeautifulSoup
from io import BytesIO
import datetime
import uuid

def extract_text_from_pdf(file):
    """Extract text from a PDF file uploaded via Streamlit."""
    file_bytes = file.getvalue()
    if not file_bytes:
        return ""

    pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in pdf_doc:
        try:
            text += page.get_text("text") + "\n"
        except Exception:
            text += page.get_text() + "\n"
    return text.strip()

def extract_text_from_docx(file):
    """Extract text from a Word document."""
    file_bytes = file.getvalue()
    doc = DocxDocument(BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_url(url):
    """Extract visible text from a webpage URL."""
    try:
        response = requests.get(url, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style"]):
            tag.extract()
        return " ".join(soup.get_text().split())
    except Exception:
        return ""

def split_into_chunks(text, chunk_size=800, overlap=200, source=""):
    """
    Split long text into overlapping chunks with enriched metadata.

    Each chunk is a langchain_core Document with metadata:
      - source: (file name or URL)
      - chunk_id: uuid4 for tracing
      - chunk_length: number of characters in chunk
      - created_at: ISO timestamp (when chunk was created)
    """
    chunks = []
    start = 0
    now_iso = datetime.datetime.now().isoformat()
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk_text = text[start:end]
        metadata = {
            "source": source or "unknown",
            "chunk_id": str(uuid.uuid4()),
            "chunk_length": len(chunk_text),
            "created_at": now_iso
        }
        chunks.append(
            Document(
                page_content=chunk_text,
                metadata=metadata
            )
        )
        start += chunk_size - overlap
    return chunks
